import streamlit as st
import requests
from bs4 import BeautifulSoup
import openai
import pandas as pd
import plotly.express as px
from docx import Document

color_profiles = {
    'Silver': {'key_characteristics': ['independence', 'self-control', 'adjustment', 'restraint', 'quiet', 'introspective', 'self-sufficient'], 'tone_and_style': ['calm', 'balanced', 'neutral', 'unbiased'], 'messaging_tips': ['clear', 'concise', 'balanced', 'neutral']},
    'Purple': {'key_characteristics': ['sensitive', 'compassionate', 'understanding', 'supportive', 'kind', 'generous', 'good listener'], 'tone_and_style': ['warm', 'empathetic', 'encouraging', 'supportive'], 'messaging_tips': ['positive reinforcement', 'encouragement', 'support', 'understanding']},
    'Pink': {'key_characteristics': ['warm', 'nurturing', 'compassionate', 'caring', 'affectionate', 'loving'], 'tone_and_style': ['gentle', 'soothing', 'nurturing', 'supportive'], 'messaging_tips': ['compassion', 'understanding', 'nurturing', 'support']},
    'Yellow': {'key_characteristics': ['happy', 'fun', 'light', 'lively', 'energetic'], 'tone_and_style': ['bright', 'uplifting', 'energetic', 'lively'], 'messaging_tips': ['positivity', 'enthusiasm', 'energy', 'brightness']},
    'Red': {'key_characteristics': ['excitement', 'strength', 'love', 'passion', 'heat', 'joy'], 'tone_and_style': ['strong', 'bold', 'passionate', 'exciting'], 'messaging_tips': ['boldness', 'excitement', 'passion', 'strength']},
    'Orange': {'key_characteristics': ['creative', 'original', 'self-expression', 'artistry', 'new ideas', 'modes of expression'], 'tone_and_style': ['exuberant', 'vivid', 'colorful', 'unrestrained', 'abstract', 'unconventional'], 'messaging_tips': ['expressive freedom', 'art for art’s sake', 'original', 'creative', 'diversity', 'imagination', 'ideation']},
    'Blue': {'key_characteristics': ['growth', 'industry leader', 'stability', 'pride', 'strength', 'influence', 'accomplishment'], 'tone_and_style': ['bold', 'confident', 'self-assured', 'proud'], 'messaging_tips': ['bold', 'confident', 'self-assured', 'proud', 'powerful']}
}

color_to_hex = {
    'Silver': '#C0C0C0',
    'Purple': '#800080',
    'Pink': '#FFC0CB',
    'Yellow': '#FFFF00',
    'Red': '#FF0000',
    'Orange': '#FFA500',
    'Blue': '#0000FF'
}

openai_api_key = st.secrets["OPENAI_API_KEY"]
openai.api_key = openai_api_key

def scrape_text(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    paragraphs = soup.find_all('p')
    text = " ".join([para.text for para in paragraphs])
    return text

def assess_content(content):
    color_guide = ""
    for color, attributes in color_profiles.items():
        color_guide += f"{color}:\\n"
        for attribute, values in attributes.items():
            color_guide += f"  {attribute}: {' '.join(values)}\\n"

    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=f"Carefully analyze the content provided and compare it with the detailed color guide below. Evaluate the content against each color’s key characteristics, tone & style, and messaging tips to determine the most fitting primary color and any supporting colors.\\n\\nContent:\\n{content}\\n\\nColor Guide:\\n{color_guide}\\n\\nBased on a detailed comparison of the content and every color profile in the color guide, identify the most aligned primary color and any supporting colors. Provide a thorough rationale explaining why each color was chosen.",
        temperature=0.5,
        max_tokens=400,
        top_p=1.0,
        frequency_penalty=0.0,
        presence_penalty=0.0
    )

    output_text = response.choices[0].text.strip()
    primary_color = "Not Identified"
    supporting_colors = "Not Identified"
    rationale = "Not Provided"
    
    lines = output_text.split('\\n')
    if lines:
        primary_color_line = lines[0].strip()
        primary_color = primary_color_line.split(":")[1].strip() if ":" in primary_color_line else primary_color_line
        if len(lines) > 1:
            supporting_colors_line = lines[1].strip()
            supporting_colors = supporting_colors_line.split(":")[1].strip() if ":" in supporting_colors_line else supporting_colors_line
            rationale = "\\n".join(lines[2:]).strip() if len(lines) > 2 else "Not Provided"

    return primary_color, supporting_colors, rationale

def create_word_document(urls_analysis):
    document = Document()
        document.add_heading('Webpage Content Color Assessor Analysis', level=1)
    for url, analysis in urls_analysis.items():
        document.add_heading(f'URL: {url}', level=2)
        document.add_heading('Primary Color:', level=3)
        document.add_paragraph(analysis['primary_color'])
        document.add_heading('Supporting Colors:', level=3)
        document.add_paragraph(', '.join(analysis['supporting_colors']))
        document.add_heading('Rationale:', level=3)
        document.add_paragraph(analysis['rationale'])
    file_path = "/mnt/data/analysis.docx"
    document.save(file_path)
    return file_path

def main():
    st.title("Webpage Content Color Assessor")
    urls_input = st.text_area("Enter up to 20 URLs (separated by commas):")
    urls = [url.strip() for url in urls_input.split(",") if url.strip()]

    if st.button("Assess Content Colors"):
        if not urls or len(urls) > 20:
            st.error("Please enter up to 20 valid URLs.")
        else:
            urls_analysis = {}
            color_count = {}
            for url in urls:
                content = scrape_text(url)
                primary_color, supporting_colors, rationale = assess_content(content)

                if f'primary_{url}' not in st.session_state:
                    st.session_state[f'primary_{url}'] = primary_color
                    st.session_state[f'supporting_{url}'] = supporting_colors.split(', ')
                    st.session_state[f'rationale_{url}'] = rationale
                
                st.write(f"**URL:** {url}")
                st.write(f"**Primary Color:** {st.session_state[f'primary_{url}']}")
                st.write(f"**Supporting Colors:** {', '.join(st.session_state[f'supporting_{url}']) if st.session_state[f'supporting_{url}'] else ''}")
                st.write(f"**Rationale:** {st.session_state[f'rationale_{url}']}")
                
                st.session_state[f'reassign_primary_{url}'] = st.selectbox(f'Reassign Primary Color for {url}', list(color_profiles.keys()), key=f'primary_{url}', value=st.session_state[f'primary_{url}'])
                st.session_state[f'reassign_supporting_{url}'] = st.multiselect(f'Reassign Supporting Colors for {url}', list(color_profiles.keys()), key=f'supporting_{url}', default=st.session_state[f'supporting_{url}'])
                st.session_state[f'retype_rationale_{url}'] = st.text_area(f'Retype Rationale for {url}', value=st.session_state[f'rationale_{url}'], key=f'rationale_{url}')
                
                urls_analysis[url] = {
                    'primary_color': st.session_state[f'primary_{url}'],
                    'supporting_colors': st.session_state[f'supporting_{url}'],
                    'rationale': st.session_state[f'rationale_{url}']
                }
                color_count[st.session_state[f'primary_{url}']] = color_count.get(st.session_state[f'primary_{url}'], 0) + 1
            
            if st.button('Update All Analyses'):
                for url in urls:
                    st.session_state[f'primary_{url}'] = st.session_state[f'reassign_primary_{url}']
                    st.session_state[f'supporting_{url}'] = st.session_state[f'reassign_supporting_{url}']
                    st.session_state[f'rationale_{url}'] = st.session_state[f'retype_rationale_{url}']

            color_count_df = pd.DataFrame(list(color_count.items()), columns=['Color', 'Count'])
            fig = px.pie(color_count_df, names='Color', values='Count', color='Color', color_discrete_map=color_to_hex, hole=0.4, width=1000, height=500)
            st.plotly_chart(fig)
            
            if st.button('Download Analysis'):
                file_path = create_word_document(urls_analysis)
                st.markdown(f'[Download Analysis]({file_path})')

if __name__ == "__main__":
    main()
