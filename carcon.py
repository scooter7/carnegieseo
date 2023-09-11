import streamlit as st
import re
import plotly.graph_objects as go
from collections import Counter
import base64
from docx import Document
from docx.shared import Inches
import io
import matplotlib.pyplot as plt

def analyze_text(text, color_keywords):
    text = text.lower()
    words = re.findall(r'\b\w+\b', text)
    color_counts = Counter()
    for color, keywords in color_keywords.items():
        color_counts[color] = sum(words.count(k.lower()) for k in keywords)
    return color_counts

def draw_donut_chart(color_counts, color_keywords):
    labels = list(color_keywords.keys())
    sizes = [color_counts.get(color, 0) for color in labels]
    colors = {label: label.lower() for label in labels}
    fig = go.Figure(data=[go.Pie(labels=labels, values=sizes, hole=.3, marker=dict(colors=[colors[label] for label in labels]))])
    return fig

def extract_examples(text, color_keywords, top_colors):
    text = text.lower()
    examples = {}
    sentences = list(set(re.split(r'[.!?]', text)))
    for color in top_colors:
        examples[color] = set()
        for keyword in color_keywords[color]:
            keyword = keyword.lower()
            for sentence in sentences:
                if keyword in sentence:
                    examples[color].add(sentence.strip() + '.')
        examples[color] = list(examples[color])[:3]
    return examples

def analyze_tone(text):
    tone_keywords = {
        "Relaxed": ["calm", "peaceful", "easygoing", "informal"],
        "Assertive": ["confident", "aggressive", "self-assured", "dogmatic"],
        "Introverted": ["calm", "solitude", "introspective", "reserved"],
        "Extroverted": ["social", "energetic", "outgoing"],
        "Conservative": ["traditional", "status quo", "orthodox"],
        "Progressive": ["reform", "liberal", "innovative"],
        "Emotive": ["emotional", "passionate", "intense"],
        "Informative": ["inform", "disclose", "instructive"]
    }
    text = text.lower()
    words = re.findall(r'\b\w+\b', text)
    tone_counts = Counter()
    for tone, keywords in tone_keywords.items():
        tone_counts[tone] = sum(words.count(k.lower()) for k in keywords)
    total_count = sum(tone_counts.values())
    tone_scores = {tone: (count / total_count) * 100 for tone, count in tone_counts.items()}
    return tone_scores

def get_word_file_download_link(file_path, filename):
    with open(file_path, "rb") as f:
        file_data = f.read()
    b64_file = base64.b64encode(file_data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_file}" download="{filename}">Download Word Report</a>'
    return href

def plot_tone_analysis(tone_scores):
    fig, ax = plt.subplots()
    ax.bar(tone_scores.keys(), tone_scores.values())
    plt.xticks(rotation=45)
    buf = io.BytesIO()
    plt.savefig(buf, format="png")
    buf.seek(0)
    return buf

def generate_word_doc(color_counts, examples, user_content, tone_scores, color_keywords):
    doc = Document()
    doc.add_heading('Color Personality Analysis', 0)
    fig = draw_donut_chart(color_counts, color_keywords)
    image_stream = io.BytesIO(fig.to_image(format="png"))
    doc.add_picture(image_stream, width=Inches(4.0))
    image_stream.close()
    tone_buf = plot_tone_analysis(tone_scores)
    doc.add_picture(tone_buf, width=Inches(4.0))
    tone_buf.close()
    for tone, score in tone_scores.items():
        doc.add_paragraph(f"{tone}: {score}%")
    for color, example_sentences in examples.items():
        doc.add_heading(f'Top Color: {color}', level=1)
        for example in example_sentences:
            doc.add_paragraph(example)
    doc.add_heading('Original Text:', level=1)
    doc.add_paragraph(user_content)
    word_file_path = "Color_Personality_Analysis_Report.docx"
    doc.save(word_file_path)
    return word_file_path

def main():
    st.title('Color Personality Analysis')
    if 'user_content' not in st.session_state:
        st.session_state.user_content = ""

    color_keywords = {
        'Red': ['Activate', 'Animate', 'Amuse', 'Captivate', 'Cheer', 'Delight', 'Encourage', 'Energize', 'Engage', 'Enjoy', 'Enliven', 'Entertain', 'Excite', 'Express', 'Inspire', 'Joke', 'Motivate', 'Play', 'Stir', 'Uplift', 'Amusing', 'Clever', 'Comedic', 'Dynamic', 'Energetic', 'Engaging', 'Enjoyable', 'Entertaining', 'Enthusiastic', 'Exciting', 'Expressive', 'Extroverted', 'Fun', 'Humorous', 'Interesting', 'Lively', 'Motivational', 'Passionate', 'Playful', 'Spirited'],
        'Silver': ['Activate', 'Campaign', 'Challenge', 'Commit', 'Confront', 'Dare', 'Defy', 'Disrupting', 'Drive', 'Excite', 'Face', 'Ignite', 'Incite', 'Influence', 'Inspire', 'Inspirit', 'Motivate', 'Move', 'Push', 'Rebel', 'Reimagine', 'Revolutionize', 'Rise', 'Spark', 'Stir', 'Fight', 'Free', 'Aggressive', 'Bold', 'Brazen', 'Committed', 'Courageous', 'Daring', 'Disruptive', 'Driven', 'Fearless', 'Free', 'Gutsy', 'Independent', 'Inspired', 'Motivated', 'Rebellious', 'Revolutionary', 'Unafraid', 'Unconventional'],
        'Blue': ['Accomplish', 'Achieve', 'Affect', 'Assert', 'Cause', 'Command', 'Determine', 'Direct', 'Dominate', 'Drive', 'Empower', 'Establish', 'Guide', 'Impact', 'Impress', 'Influence', 'Inspire', 'Lead', 'Outpace', 'Outshine', 'Realize', 'Shape', 'Succeed', 'Transform', 'Win', 'Accomplished', 'Assertive', 'Authoritative', 'Commanding', 'Confident', 'Decisive', 'Distinguished', 'Dominant', 'Elite', 'Eminent', 'Established', 'Exceptional', 'Expert', 'First-class', 'First-rate', 'Impressive', 'Influential', 'Leading', 'Magnetic', 'Managerial', 'Masterful', 'Noble', 'Premier', 'Prestigious', 'Prominent', 'Proud', 'Strong'],
        'Yellow': ['Accelerate', 'Advance', 'Change', 'Conceive', 'Create', 'Engineer', 'Envision', 'Experiment', 'Dream', 'Ignite', 'Illuminate', 'Imagine', 'Innovate', 'Inspire', 'Invent', 'Pioneer', 'Progress', 'Shape', 'Spark', 'Solve', 'Transform', 'Unleash', 'Unlock', 'Advanced', 'Brilliant', 'Conceptual', 'Enterprising', 'Expert', 'Extraordinary', 'Forward-looking', 'Forward-thinking', 'Fresh', 'Future-minded', 'Future-thinking', 'Ingenious', 'Intelligent', 'Inventive', 'Leading-edge', 'Luminous', 'New', 'Pioneering', 'Reforming', 'Rising', 'Transformative', 'Visionary', 'World-changing', 'World-class'],
        'Green': ['Analyze', 'Discover', 'Examine', 'Expand', 'Explore', 'Extend', 'Inquire', 'Journey', 'Launch', 'Move', 'Pioneer', 'Pursue', 'Question', 'Reach', 'Search', 'Uncover', 'Venture', 'Wonder', 'Adventurous', 'Analytical', 'Curious', 'Discerning', 'Experiential', 'Exploratory', 'Fearless', 'Inquisitive', 'Intriguing', 'Investigative', 'Journeying', 'Mysterious', 'Philosophical', 'Pioneering', 'Questioning', 'Unbound', 'Unexpected'],
        'Purple': ['Accommodate', 'Assist', 'Befriend', 'Care', 'Collaborate', 'Connect', 'Embrace', 'Empower', 'Encourage', 'Foster', 'Give', 'Help', 'Nourish', 'Nurture', 'Promote', 'Protect', 'Provide', 'Serve', 'Share', 'Shepherd', 'Steward', 'Tend', 'Uplift', 'Value', 'Welcome', 'Affectionate', 'Attentive', 'Beneficial', 'Benevolent', 'Big-hearted', 'Caring', 'Charitable', 'Compassionate', 'Considerate', 'Encouraging', 'Friendly', 'Generous', 'Gentle', 'Helpful', 'Hospitable', 'Inclusive', 'Kind-hearted', 'Merciful', 'Missional', 'Neighborly', 'Nurturing', 'Protective', 'Responsible', 'Selfless', 'Supportive', 'Sympathetic', 'Thoughtful', 'Uplifting', 'Vocational', 'Warm'],
        'Maroon': ['Accomplish', 'Achieve', 'Build', 'Challenge', 'Commit', 'Compete', 'Contend', 'Dedicate', 'Defend', 'Devote', 'Drive', 'Endeavor', 'Entrust', 'Endure', 'Fight', 'Grapple', 'Grow', 'Improve', 'Increase', 'Overcome', 'Persevere', 'Persist', 'Press on', 'Pursue', 'Resolve', 'Tackle', 'Ambitious', 'Brave', 'Committed', 'Competitive', 'Consistent', 'Constant', 'Continuous', 'Courageous', 'Dedicated', 'Determined', 'Earnest', 'Industrious', 'Loyal', 'Persevering', 'Persistent', 'Proud', 'Purposeful', 'Relentless', 'Reliable', 'Resilient', 'Resolute', 'Steadfast', 'Strong', 'Tenacious', 'Tireless', 'Tough'],
        'Orange': ['Compose', 'Conceptualize', 'Conceive', 'Craft', 'Create', 'Design', 'Dream', 'Envision', 'Express', 'Fashion', 'Form', 'Imagine', 'Interpret', 'Make', 'Originate', 'Paint', 'Perform', 'Portray', 'Realize', 'Shape', 'Abstract', 'Artistic', 'Avant-garde', 'Colorful', 'Conceptual', 'Contemporary', 'Creative', 'Decorative', 'Eccentric', 'Eclectic', 'Evocative', 'Expressive', 'Imaginative', 'Interpretive', 'Offbeat', 'One-of-a-kind', 'Original', 'Uncommon', 'Unconventional', 'Unexpected', 'Unique', 'Vibrant', 'Whimsical'],
        'Pink': ['Arise', 'Aspire', 'Detail', 'Dream', 'Elevate', 'Enchant', 'Enrich', 'Envision', 'Exceed', 'Excel', 'Experience', 'Improve', 'Idealize', 'Imagine', 'Inspire', 'Perfect', 'Poise', 'Polish', 'Prepare', 'Refine', 'Uplift', 'Affectionate', 'Admirable', 'Age-less', 'Beautiful', 'Classic', 'Desirable', 'Detailed', 'Dreamy', 'Elegant', 'Enchanting', 'Enriching', 'Ethereal', 'Excellent', 'Exceptional', 'Experiential', 'Exquisite', 'Glamorous', 'Graceful', 'Idealistic', 'Inspiring', 'Lofty', 'Mysterious', 'Ordered', 'Perfect', 'Poised', 'Polished', 'Pristine', 'Pure', 'Refined', 'Romantic', 'Sophisticated', 'Spiritual', 'Timeless', 'Traditional', 'Virtuous', 'Visionary']
    }

    user_content = st.text_area('Paste your content here:', value=st.session_state.user_content)
    st.session_state.user_content = user_content
    
    if st.button('Analyze'):
        color_counts = analyze_text(user_content, color_keywords)
        total_counts = sum(color_counts.values())
        if total_counts == 0:
            st.write('No relevant keywords found.')
            return
        fig = draw_donut_chart(color_counts, color_keywords)
        st.plotly_chart(fig)
        sorted_colors = sorted(color_counts.items(), key=lambda x: x[1], reverse=True)
        top_colors = [color for color, _ in sorted_colors[:3]]
        examples = extract_examples(user_content, color_keywords, top_colors)
        for color in top_colors:
            st.write(f'Examples for {color}:')
            st.write(', '.join(examples[color]))
        tone_scores = analyze_tone(user_content)
        st.subheader("Tone Analysis")
        st.write("The text exhibits the following tones:")
        st.bar_chart(tone_scores)
        word_file_path = generate_word_doc(color_counts, examples, user_content, tone_scores, color_keywords)
        download_link = get_word_file_download_link(word_file_path, "Color_Personality_Analysis_Report.docx")
        st.markdown(download_link, unsafe_allow_html=True)

    st.subheader("Revision Field")
    revision_input = st.text_area("Paste a sentence here for revision:")
    revised_color = st.selectbox("Select the revised color:", list(color_keywords.keys()))
    
    if st.button("Apply Revision"):
        if revision_input:
            pattern = re.escape(revision_input.strip())
            if pattern in user_content:
                revised_sentence = f"{revision_input.strip()} ({revised_color})"
                st.session_state.user_content = user_content.replace(revision_input, revised_sentence)
                st.success(f"Sentence revised to include '{revised_color}'.")
            else:
                st.error("The sentence was not found in the original content.")

    st.text_area('Updated content:', value=st.session_state.user_content)

if __name__ == '__main__':
    main()
