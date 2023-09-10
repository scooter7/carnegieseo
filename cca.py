import streamlit as st
import re
from collections import Counter
import base64
from docx import Document
from docx.shared import Inches
import plotly.graph_objects as go
import openai

def analyze_text(text, color_keywords):
    text = text.lower()
    words = re.findall(r'\b\w+\b', text)
    color_counts = Counter()
    for color, keywords in color_keywords.items():
        color_counts[color] = sum(words.count(k.lower()) for k in keywords)
    return color_counts

def draw_donut_chart(labels, sizes):
    colors = labels
    fig = go.Figure(data=[go.Pie(labels=labels, values=sizes, hole=.3, marker=dict(colors=colors))])
    fig.write_image("donut_chart.png")
    return fig

def draw_quadrant_chart(tone_analysis, title, x_axis_labels, y_axis_labels):
    fig = go.Figure()
    x_axis = list(tone_analysis.keys())
    y_axis = list(tone_analysis.values())
    
    labels = [f'{k}: {v}' for k, v in tone_analysis.items()]
    
    fig.add_trace(go.Scatter(x=x_axis, y=y_axis, text=labels, mode='markers+text', name='Tone'))
    
    fig.add_shape(type='line', x0=5, x1=5, y0=0, y1=10, line=dict(color='Grey', width=1, dash='dash'))
    fig.add_shape(type='line', x0=0, x1=10, y0=5, y1=5, line=dict(color='Grey', width=1, dash='dash'))
    
    fig.update_xaxes(range=[0, 10], tickvals=list(range(0, 11)), ticktext=['', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10'])
    fig.update_yaxes(range=[0, 10], tickvals=list(range(0, 11)), ticktext=['', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10'])
    
    annotations = []
    annotations.append(dict(x=5, y=9, xref='x', yref='y', text=x_axis_labels[0], showarrow=False, ha='center'))
    annotations.append(dict(x=5, y=1, xref='x', yref='y', text=x_axis_labels[1], showarrow=False, ha='center'))
    annotations.append(dict(x=1, y=5, xref='x', yref='y', text=y_axis_labels[0], showarrow=False, ha='center'))
    annotations.append(dict(x=9, y=5, xref='x', yref='y', text=y_axis_labels[1], showarrow=False, ha='center'))
    
    fig.update_layout(title=title, annotations=annotations)
    fig.write_image(title + ".png")
    return fig

def extract_examples(text, color_keywords, top_colors):
    text = text.lower()
    examples = {}
    sentences = list(set(re.split(r'[.!?]', text)))
    for color in top_colors:
        examples[color] = set()
        for keyword in color_keywords[color]:
            keyword = keyword.lower()
            for sentence in sentences:
                if keyword in sentence:
                    examples[color].add(sentence.strip() + '.')
        examples[color] = list(examples[color])[:3]
    return examples

def analyze_with_gpt3(text, api_key):
    openai.api_key = api_key
    response = openai.Completion.create(
        engine='text-davinci-002',
        prompt=text,
        max_tokens=50,
        temperature=0.5
    )
    
    # Debugging: Log GPT-3 Response
    print("GPT-3 Response:", response.choices[0].text.strip())
    
    tone_scores = {}
    score_pairs = response.choices[0].text.strip().split(', ')
    
    for pair in score_pairs:
        # Error Handling: Skip lines that can't be split into key and value
        if ': ' not in pair:
            continue
        
        key, value = pair.split(': ')
        tone_scores[key.strip()] = int(value.strip())
    
    return tone_scores

def generate_word_doc(top_colors, examples, user_content, general_analysis, tone_scores, new_tone_scores):
    doc = Document()
    doc.add_heading('Color Personality Analysis', 0)
    doc.add_picture('donut_chart.png', width=Inches(4.0))
    doc.add_heading('Top Colors', level=1)
    for color in top_colors:
        doc.add_paragraph(f'Top Color: {color}')
        for example in examples[color]:
            doc.add_paragraph(example, style='ListBullet')
    doc.add_heading('Original Content', level=1)
    doc.add_paragraph(user_content)
    doc.add_heading('GPT-3 Analysis', level=1)
    doc.add_paragraph(general_analysis)
    doc.add_heading('Tone Analysis', level=1)
    doc.add_paragraph(', '.join(f'{k}: {v}' for k, v in tone_scores.items()))
    doc.add_picture('Tone Quadrant Chart.png', width=Inches(4.0))
    doc.add_heading('Additional Tone Analysis', level=1)
    doc.add_paragraph(', '.join(f'{k}: {v}' for k, v in new_tone_scores.items()))
    doc.add_picture('Additional Tone Quadrant Chart.png', width=Inches(4.0))
    word_file_path = 'report.docx'
    doc.save(word_file_path)
    return word_file_path

def download_file(file_path):
    with open(file_path, 'rb') as f:
        file_data = f.read()
    b64_file = base64.b64encode(file_data).decode('utf-8')
    st.markdown(f'<a href="data:application/octet-stream;base64,{b64_file}" download="{file_path}">Download {file_path}</a>', unsafe_allow_html=True)

def main():
    st.title('Color Personality Analysis')
    if 'OPENAI_API_KEY' not in st.secrets:
        st.error('Please set the OPENAI_API_KEY secret on the Streamlit dashboard.')
        return

    openai_api_key = st.secrets['OPENAI_API_KEY']
    color_keywords = {
        'Red': ['Activate', 'Animate', 'Amuse', 'Captivate', 'Cheer', 'Delight', 'Encourage', 'Energize', 'Engage', 'Enjoy', 'Enliven', 'Entertain', 'Excite', 'Express', 'Inspire', 'Joke', 'Motivate', 'Play', 'Stir', 'Uplift', 'Amusing', 'Clever', 'Comedic', 'Dynamic', 'Energetic', 'Engaging', 'Enjoyable', 'Entertaining', 'Enthusiastic', 'Exciting', 'Expressive', 'Extroverted', 'Fun', 'Humorous', 'Interesting', 'Lively', 'Motivational', 'Passionate', 'Playful', 'Spirited'],
        'Silver': ['Activate', 'Campaign', 'Challenge', 'Commit', 'Confront', 'Dare', 'Defy', 'Disrupting', 'Drive', 'Excite', 'Face', 'Ignite', 'Incite', 'Influence', 'Inspire', 'Inspirit', 'Motivate', 'Move', 'Push', 'Rebel', 'Reimagine', 'Revolutionize', 'Rise', 'Spark', 'Stir', 'Fight', 'Free', 'Aggressive', 'Bold', 'Brazen', 'Committed', 'Courageous', 'Daring', 'Disruptive', 'Driven', 'Fearless', 'Free', 'Gutsy', 'Independent', 'Inspired', 'Motivated', 'Rebellious', 'Revolutionary', 'Unafraid', 'Unconventional'],
        'Blue': ['Accomplish', 'Achieve', 'Affect', 'Assert', 'Cause', 'Command', 'Determine', 'Direct', 'Dominate', 'Drive', 'Empower', 'Establish', 'Guide', 'Impact', 'Impress', 'Influence', 'Inspire', 'Lead', 'Outpace', 'Outshine', 'Realize', 'Shape', 'Succeed', 'Transform', 'Win', 'Accomplished', 'Assertive', 'Authoritative', 'Commanding', 'Confident', 'Decisive', 'Distinguished', 'Dominant', 'Elite', 'Eminent', 'Established', 'Exceptional', 'Expert', 'First-class', 'First-rate', 'Impressive', 'Influential', 'Leading', 'Magnetic', 'Managerial', 'Masterful', 'Noble', 'Premier', 'Prestigious', 'Prominent', 'Proud', 'Strong'],
        'Yellow': ['Accelerate', 'Advance', 'Change', 'Conceive', 'Create', 'Engineer', 'Envision', 'Experiment', 'Dream', 'Ignite', 'Illuminate', 'Imagine', 'Innovate', 'Inspire', 'Invent', 'Pioneer', 'Progress', 'Shape', 'Spark', 'Solve', 'Transform', 'Unleash', 'Unlock', 'Advanced', 'Brilliant', 'Conceptual', 'Enterprising', 'Expert', 'Extraordinary', 'Forward-looking', 'Forward-thinking', 'Fresh', 'Future-minded', 'Future-thinking', 'Ingenious', 'Intelligent', 'Inventive', 'Leading-edge', 'Luminous', 'New', 'Pioneering', 'Reforming', 'Rising', 'Transformative', 'Visionary', 'World-changing', 'World-class'],
        'Green': ['Analyze', 'Discover', 'Examine', 'Expand', 'Explore', 'Extend', 'Inquire', 'Journey', 'Launch', 'Move', 'Pioneer', 'Pursue', 'Question', 'Reach', 'Search', 'Uncover', 'Venture', 'Wonder', 'Adventurous', 'Analytical', 'Curious', 'Discerning', 'Experiential', 'Exploratory', 'Fearless', 'Inquisitive', 'Intriguing', 'Investigative', 'Journeying', 'Mysterious', 'Philosophical', 'Pioneering', 'Questioning', 'Unbound', 'Unexpected'],
        'Purple': ['Accommodate', 'Assist', 'Befriend', 'Care', 'Collaborate', 'Connect', 'Embrace', 'Empower', 'Encourage', 'Foster', 'Give', 'Help', 'Nourish', 'Nurture', 'Promote', 'Protect', 'Provide', 'Serve', 'Share', 'Shepherd', 'Steward', 'Tend', 'Uplift', 'Value', 'Welcome', 'Affectionate', 'Attentive', 'Beneficial', 'Benevolent', 'Big-hearted', 'Caring', 'Charitable', 'Compassionate', 'Considerate', 'Encouraging', 'Friendly', 'Generous', 'Gentle', 'Helpful', 'Hospitable', 'Inclusive', 'Kind-hearted', 'Merciful', 'Missional', 'Neighborly', 'Nurturing', 'Protective', 'Responsible', 'Selfless', 'Supportive', 'Sympathetic', 'Thoughtful', 'Uplifting', 'Vocational', 'Warm'],
        'Maroon': ['Accomplish', 'Achieve', 'Build', 'Challenge', 'Commit', 'Compete', 'Contend', 'Dedicate', 'Defend', 'Devote', 'Drive', 'Endeavor', 'Entrust', 'Endure', 'Fight', 'Grapple', 'Grow', 'Improve', 'Increase', 'Overcome', 'Persevere', 'Persist', 'Press on', 'Pursue', 'Resolve', 'Tackle', 'Ambitious', 'Brave', 'Committed', 'Competitive', 'Consistent', 'Constant', 'Continuous', 'Courageous', 'Dedicated', 'Determined', 'Earnest', 'Industrious', 'Loyal', 'Persevering', 'Persistent', 'Proud', 'Purposeful', 'Relentless', 'Reliable', 'Resilient', 'Resolute', 'Steadfast', 'Strong', 'Tenacious', 'Tireless', 'Tough'],
        'Orange': ['Compose', 'Conceptualize', 'Conceive', 'Craft', 'Create', 'Design', 'Dream', 'Envision', 'Express', 'Fashion', 'Form', 'Imagine', 'Interpret', 'Make', 'Originate', 'Paint', 'Perform', 'Portray', 'Realize', 'Shape', 'Abstract', 'Artistic', 'Avant-garde', 'Colorful', 'Conceptual', 'Contemporary', 'Creative', 'Decorative', 'Eccentric', 'Eclectic', 'Evocative', 'Expressive', 'Imaginative', 'Interpretive', 'Offbeat', 'One-of-a-kind', 'Original', 'Uncommon', 'Unconventional', 'Unexpected', 'Unique', 'Vibrant', 'Whimsical'],
        'Pink': ['Arise', 'Aspire', 'Detail', 'Dream', 'Elevate', 'Enchant', 'Enrich', 'Envision', 'Exceed', 'Excel', 'Experience', 'Improve', 'Idealize', 'Imagine', 'Inspire', 'Perfect', 'Poise', 'Polish', 'Prepare', 'Refine', 'Uplift', 'Affectionate', 'Admirable', 'Age-less', 'Beautiful', 'Classic', 'Desirable', 'Detailed', 'Dreamy', 'Elegant', 'Enchanting', 'Enriching', 'Ethereal', 'Excellent', 'Exceptional', 'Experiential', 'Exquisite', 'Glamorous', 'Graceful', 'Idealistic', 'Inspiring', 'Lofty', 'Mysterious', 'Ordered', 'Perfect', 'Poised', 'Polished', 'Pristine', 'Pure', 'Refined', 'Romantic', 'Sophisticated', 'Spiritual', 'Timeless', 'Traditional', 'Virtuous', 'Visionary']
    }

    user_content = st.text_area('Paste your content here:')

    if st.button('Analyze'):
        color_counts = analyze_text(user_content, color_keywords)
        total_counts = sum(color_counts.values())
        if total_counts == 0:
            st.write('No relevant keywords found.')
            return

        sorted_colors = sorted(color_counts.items(), key=lambda x: x[1], reverse=True)
        top_colors = [color for color, _ in sorted_colors[:3]]
        labels = [k for k, v in color_counts.items() if v > 0]
        sizes = [v for v in color_counts.values() if v > 0]
        fig = draw_donut_chart(labels, sizes)
        st.plotly_chart(fig)
        examples = extract_examples(user_content, color_keywords, top_colors)
        for color in top_colors:
            st.write(f'Examples for {color}:')
            st.write(', '.join(examples[color]))

        general_analysis = analyze_with_gpt3(user_content, openai_api_key)
        st.write('GPT-3 Analysis:')
        st.write(general_analysis)

        # Analyze tone with GPT-3
        tone_analysis_prompt = 'Assess the text for tone. Provide scores for the following four traits: relaxed, assertive, introverted, extroverted.'
        tone_scores = analyze_with_gpt3(user_content, openai_api_key)

        # Analyze additional tone with GPT-3
        additional_tone_prompt = 'Assess the text for additional tone. Provide scores for the following four traits: conservative, progressive, emotive, informative.'
        new_tone_scores = analyze_with_gpt3(user_content, openai_api_key)

        tone_scores = {k: int(v) for k, v in tone_scores.items()}
        new_tone_scores = {k: int(v) for k, v in new_tone_scores.items()}

        fig1 = draw_quadrant_chart(tone_scores, 'Tone Quadrant Chart', ['Relaxed', 'Assertive'], ['Extroverted', 'Introverted'])
        fig2 = draw_quadrant_chart(new_tone_scores, 'Additional Tone Quadrant Chart', ['Conservative', 'Progressive'], ['Emotive', 'Informative'])

        st.plotly_chart(fig1)
        st.plotly_chart(fig2)

        word_file_path = generate_word_doc(top_colors, examples, user_content, general_analysis, tone_scores, new_tone_scores)
        download_file(word_file_path)

if __name__ == '__main__':
    main()
