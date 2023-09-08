import streamlit as st
import re
import plotly.express as px
from collections import Counter
import base64
from docx import Document
from docx.shared import Inches
import openai

def analyze_text(text, color_keywords):
    text = text.lower()
    words = re.findall(r'\b\w+\b', text)
    color_counts = Counter()
    for color, keywords in color_keywords.items():
        color_counts[color] = sum(words.count(k.lower()) for k in keywords)
    return color_counts

def draw_column_chart(labels, sizes):
    fig = px.bar(x=labels, y=sizes)
    fig.write_image("chart.png")
    return fig

def extract_examples(text, color_keywords, top_colors):
    text = text.lower()
    examples = {}
    sentences = list(set(re.split(r'[.!?]', text)))
    for color in top_colors:
        examples[color] = set()
        for keyword in color_keywords[color]:
            keyword = keyword.lower()
            for sentence in sentences:
                if keyword in sentence:
                    examples[color].add(sentence.strip() + '.')
        examples[color] = list(examples[color])[:3]
    return examples

def analyze_with_gpt3(text, api_key):
    openai.api_key = api_key
    prompt = f"Please analyze the following text and identify who would likely find it compelling:\n\n{text}"
    response = openai.Completion.create(engine="text-davinci-002", prompt=prompt, max_tokens=100)
    return response.choices[0].text.strip()

def generate_word_doc(top_colors, examples, user_content, gpt3_analysis):
    doc = Document()
    doc.add_heading('Color Personality Analysis', 0)
    doc.add_picture('chart.png', width=Inches(4.0))
    for color in top_colors:
        doc.add_heading(f'Top Color: {color}', level=1)
        for example in examples[color]:
            doc.add_paragraph(example)
    doc.add_heading('Original Text:', level=1)
    doc.add_paragraph(user_content)
    doc.add_heading('GPT-3 Analysis:', level=1)
    doc.add_paragraph(gpt3_analysis)
    word_file_path = "report.docx"
    doc.save(word_file_path)
    return word_file_path

def download_file(file_path):
    with open(file_path, "rb") as f:
        file_data = f.read()
    b64_file = base64.b64encode(file_data).decode("utf-8")
    href = f'<a href="data:application/octet-stream;base64,{b64_file}" download="report.docx">Download Word Report</a>'
    st.markdown(href, unsafe_allow_html=True)

def main():
    st.title("Color Personality Analysis")
    color_keywords = {
        'Red': ['Activate', 'Animate', 'Amuse', 'Captivate', 'Cheer', 'Delight', 'Encourage', 'Energize', 'Engage', 'Enjoy', 'Enliven', 'Entertain', 'Excite', 'Express', 'Inspire', 'Joke', 'Motivate', 'Play', 'Stir', 'Uplift', 'Amusing', 'Clever', 'Comedic', 'Dynamic', 'Energetic', 'Engaging', 'Enjoyable', 'Entertaining', 'Enthusiastic', 'Exciting', 'Expressive', 'Extroverted', 'Fun', 'Humorous', 'Interesting', 'Lively', 'Motivational', 'Passionate', 'Playful', 'Spirited'],
        'Silver': ['Activate', 'Campaign', 'Challenge', 'Commit', 'Confront', 'Dare', 'Defy', 'Disrupting', 'Drive', 'Excite', 'Face', 'Ignite', 'Incite', 'Influence', 'Inspire', 'Inspirit', 'Motivate', 'Move', 'Push', 'Rebel', 'Reimagine', 'Revolutionize', 'Rise', 'Spark', 'Stir', 'Fight', 'Free', 'Aggressive', 'Bold', 'Brazen', 'Committed', 'Courageous', 'Daring', 'Disruptive', 'Driven', 'Fearless', 'Free', 'Gutsy', 'Independent', 'Inspired', 'Motivated', 'Rebellious', 'Revolutionary', 'Unafraid', 'Unconventional'],
        'Blue': ['Accomplish', 'Achieve', 'Affect', 'Assert', 'Cause', 'Command', 'Determine', 'Direct', 'Dominate', 'Drive', 'Empower', 'Establish', 'Guide', 'Impact', 'Impress', 'Influence', 'Inspire', 'Lead', 'Outpace', 'Outshine', 'Realize', 'Shape', 'Succeed', 'Transform', 'Win', 'Accomplished', 'Assertive', 'Authoritative', 'Commanding', 'Confident', 'Decisive', 'Distinguished', 'Dominant', 'Elite', 'Eminent', 'Established', 'Exceptional', 'Expert', 'First-class', 'First-rate', 'Impressive', 'Influential', 'Leading', 'Magnetic', 'Managerial', 'Masterful', 'Noble', 'Premier', 'Prestigious', 'Prominent', 'Proud', 'Strong'],
        'Yellow': ['Accelerate', 'Advance', 'Change', 'Conceive', 'Create', 'Engineer', 'Envision', 'Experiment', 'Dream', 'Ignite', 'Illuminate', 'Imagine', 'Innovate', 'Inspire', 'Invent', 'Pioneer', 'Progress', 'Shape', 'Spark', 'Solve', 'Transform', 'Unleash', 'Unlock', 'Advanced', 'Brilliant', 'Conceptual', 'Enterprising', 'Expert', 'Extraordinary', 'Forward-looking', 'Forward-thinking', 'Fresh', 'Future-minded', 'Future-thinking', 'Ingenious', 'Intelligent', 'Inventive', 'Leading-edge', 'Luminous', 'New', 'Pioneering', 'Reforming', 'Rising', 'Transformative', 'Visionary', 'World-changing', 'World-class'],
        'Green': ['Analyze', 'Discover', 'Examine', 'Expand', 'Explore', 'Extend', 'Inquire', 'Journey', 'Launch', 'Move', 'Pioneer', 'Pursue', 'Question', 'Reach', 'Search', 'Uncover', 'Venture', 'Wonder', 'Adventurous', 'Analytical', 'Curious', 'Discerning', 'Experiential', 'Exploratory', 'Fearless', 'Inquisitive', 'Intriguing', 'Investigative', 'Journeying', 'Mysterious', 'Philosophical', 'Pioneering', 'Questioning', 'Unbound', 'Unexpected'],
        'Purple': ['Accommodate', 'Assist', 'Befriend', 'Care', 'Collaborate', 'Connect', 'Embrace', 'Empower', 'Encourage', 'Foster', 'Give', 'Help', 'Nourish', 'Nurture', 'Promote', 'Protect', 'Provide', 'Serve', 'Share', 'Shepherd', 'Steward', 'Tend', 'Uplift', 'Value', 'Welcome', 'Affectionate', 'Attentive', 'Beneficial', 'Benevolent', 'Big-hearted', 'Caring', 'Charitable', 'Compassionate', 'Considerate', 'Encouraging', 'Friendly', 'Generous', 'Gentle', 'Helpful', 'Hospitable', 'Inclusive', 'Kind-hearted', 'Merciful', 'Missional', 'Neighborly', 'Nurturing', 'Protective', 'Responsible', 'Selfless', 'Supportive', 'Sympathetic', 'Thoughtful', 'Uplifting', 'Vocational', 'Warm'],
        'Maroon': ['Accomplish', 'Achieve', 'Build', 'Challenge', 'Commit', 'Compete', 'Contend', 'Dedicate', 'Defend', 'Devote', 'Drive', 'Endeavor', 'Entrust', 'Endure', 'Fight', 'Grapple', 'Grow', 'Improve', 'Increase', 'Overcome', 'Persevere', 'Persist', 'Press on', 'Pursue', 'Resolve', 'Tackle', 'Ambitious', 'Brave', 'Committed', 'Competitive', 'Consistent', 'Constant', 'Continuous', 'Courageous', 'Dedicated', 'Determined', 'Earnest', 'Industrious', 'Loyal', 'Persevering', 'Persistent', 'Proud', 'Purposeful', 'Relentless', 'Reliable', 'Resilient', 'Resolute', 'Steadfast', 'Strong', 'Tenacious', 'Tireless', 'Tough'],
        'Orange': ['Compose', 'Conceptualize', 'Conceive', 'Craft', 'Create', 'Design', 'Dream', 'Envision', 'Express', 'Fashion', 'Form', 'Imagine', 'Interpret', 'Make', 'Originate', 'Paint', 'Perform', 'Portray', 'Realize', 'Shape', 'Abstract', 'Artistic', 'Avant-garde', 'Colorful', 'Conceptual', 'Contemporary', 'Creative', 'Decorative', 'Eccentric', 'Eclectic', 'Evocative', 'Expressive', 'Imaginative', 'Interpretive', 'Offbeat', 'One-of-a-kind', 'Original', 'Uncommon', 'Unconventional', 'Unexpected', 'Unique', 'Vibrant', 'Whimsical'],
        'Pink': ['Arise', 'Aspire', 'Detail', 'Dream', 'Elevate', 'Enchant', 'Enrich', 'Envision', 'Exceed', 'Excel', 'Experience', 'Improve', 'Idealize', 'Imagine', 'Inspire', 'Perfect', 'Poise', 'Polish', 'Prepare', 'Refine', 'Uplift', 'Affectionate', 'Admirable', 'Age-less', 'Beautiful', 'Classic', 'Desirable', 'Detailed', 'Dreamy', 'Elegant', 'Enchanting', 'Enriching', 'Ethereal', 'Excellent', 'Exceptional', 'Experiential', 'Exquisite', 'Glamorous', 'Graceful', 'Idealistic', 'Inspiring', 'Lofty', 'Mysterious', 'Ordered', 'Perfect', 'Poised', 'Polished', 'Pristine', 'Pure', 'Refined', 'Romantic', 'Sophisticated', 'Spiritual', 'Timeless', 'Traditional', 'Virtuous', 'Visionary']
    }
    user_content = st.text_area("Paste your content here:")
    if "OPENAI_API_KEY" not in st.secrets:
        st.error("Please set the OPENAI_API_KEY secret on the Streamlit dashboard.")
        return
    openai_api_key = st.secrets["OPENAI_API_KEY"]
    if st.button('Analyze'):
        st.write("Original Text:")
        st.write(user_content)
        color_counts = analyze_text(user_content, color_keywords)
        total_counts = sum(color_counts.values())
        if total_counts == 0:
            st.write("No relevant keywords found.")
            return
        sorted_colors = sorted(color_counts.items(), key=lambda x: x[1], reverse=True)
        top_colors = [color for color, _ in sorted_colors[:3]]
        labels = [k for k, v in color_counts.items() if v > 0]
        sizes = [v for v in color_counts.values() if v > 0]
        fig = draw_column_chart(labels, sizes)
        st.plotly_chart(fig)
        examples = extract_examples(user_content, color_keywords, top_colors)
        for color in top_colors:
            st.write(f"Examples for {color}:")
            st.write(", ".join(examples[color]))
        gpt3_analysis = analyze_with_gpt3(user_content, openai_api_key)
        st.write("GPT-3 Analysis:")
        st.write(gpt3_analysis)
        word_file_path = generate_word_doc(top_colors, examples, user_content, gpt3_analysis)
        download_file(word_file_path)

main()
