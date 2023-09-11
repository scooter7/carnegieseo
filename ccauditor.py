import streamlit as st
import re
import plotly.graph_objects as go
from collections import Counter
from docx import Document
from docx.shared import Inches
import io
import base64
import openai

def analyze_text(text, color_keywords):
    text = text.lower()
    words = re.findall(r'\b\w+\b', text)
    color_counts = Counter()
    for color, keywords in color_keywords.items():
        color_counts[color] = sum(words.count(k.lower()) for k in keywords)
    return color_counts

def draw_donut_chart(color_counts):
    labels = list(color_counts.keys())
    sizes = [color_counts.get(color, 0) for color in labels]
    fig = go.Figure(data=[go.Pie(labels=labels, values=sizes, hole=.3, marker=dict(colors=labels))])
    return fig

def analyze_tone_with_gpt3(text, api_key):
    openai.api_key = api_key
    prompt = f"""
    Please provide a nuanced analysis of the following text, assigning a level to indicate the extent to which the text embodies each of the following tones:
    - Relaxed
    - Assertive
    - Introverted
    - Extroverted
    - Conservative
    - Progressive
    - Emotive
    - Informative
    Text to Analyze:
    {text}
    """
    response = openai.Completion.create(engine="text-davinci-002", prompt=prompt, max_tokens=100)
    gpt3_output = response.choices[0].text.strip().split('\n')
    tone_scores = {}
    for line in gpt3_output:
        if ":" in line:
            tone, score = line.split(":")
            if score.strip():
                tone_scores[tone.strip()] = float(score.strip())
    return tone_scores

def generate_word_doc(color_counts, user_content, tone_scores, initial_fig, tone_fig, updated_fig):
    doc = Document()
    doc.add_heading('Color Personality Analysis', 0)
    image_stream = io.BytesIO(initial_fig.to_image(format="png"))
    doc.add_heading('Initial Donut Chart:', level=1)
    doc.add_picture(image_stream, width=Inches(4.0))
    image_stream.close()
    image_stream = io.BytesIO(tone_fig.to_image(format="png"))
    doc.add_heading('Tone Analysis:', level=1)
    doc.add_picture(image_stream, width=Inches(4.0))
    image_stream.close()
    image_stream = io.BytesIO(updated_fig.to_image(format="png"))
    doc.add_heading('Updated Donut Chart:', level=1)
    doc.add_picture(image_stream, width=Inches(4.0))
    image_stream.close()
    doc.add_heading('Tone Scores:', level=1)
    for tone, score in tone_scores.items():
        doc.add_paragraph(f"{tone}: {score}")
    doc.add_heading('Scored Sentences:', level=1)
    for sentence, colors in st.session_state.sentence_to_colors.items():
        doc.add_paragraph(f"{sentence}: {', '.join(colors)}")
    doc.add_heading('Original Text:', level=1)
    doc.add_paragraph(user_content)
    word_file_path = "Color_Personality_Analysis_Report.docx"
    doc.save(word_file_path)
    return word_file_path

def get_word_file_download_link(file_path, filename):
    with open(file_path, "rb") as f:
        file_data = f.read()
    b64_file = base64.b64encode(file_data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_file}" download="{filename}">Download Word Report</a>'
    return href

def main():
    st.title('Color Personality Analysis')
    
    if "OPENAI_API_KEY" not in st.secrets:
        st.error("Please set the OPENAI_API_KEY secret on the Streamlit dashboard.")
        return
    
    openai_api_key = st.secrets["OPENAI_API_KEY"]
    
    if 'init_done' not in st.session_state:
        st.session_state.init_done = False
        st.session_state.tone_scores = {}
        st.session_state.sentence_to_colors = {}
        st.session_state.updated_color_counts = Counter()
        st.session_state.initial_fig = None
        st.session_state.tone_fig = None
        st.session_state.updated_fig = None
    
    color_keywords = {
        'Red': ['Activate', 'Animate', 'Amuse', 'Captivate', 'Cheer', 'Delight', 'Encourage', 'Energize', 'Engage', 'Enjoy', 'Enliven', 'Entertain', 'Excite', 'Express', 'Inspire', 'Joke', 'Motivate', 'Play', 'Stir', 'Uplift', 'Amusing', 'Clever', 'Comedic', 'Dynamic', 'Energetic', 'Engaging', 'Enjoyable', 'Entertaining', 'Enthusiastic', 'Exciting', 'Expressive', 'Extroverted', 'Fun', 'Humorous', 'Interesting', 'Lively', 'Motivational', 'Passionate', 'Playful', 'Spirited'],
        'Silver': ['Activate', 'Campaign', 'Challenge', 'Commit', 'Confront', 'Dare', 'Defy', 'Disrupting', 'Drive', 'Excite', 'Face', 'Ignite', 'Incite', 'Influence', 'Inspire', 'Inspirit', 'Motivate', 'Move', 'Push', 'Rebel', 'Reimagine', 'Revolutionize', 'Rise', 'Spark', 'Stir', 'Fight', 'Free', 'Aggressive', 'Bold', 'Brazen', 'Committed', 'Courageous', 'Daring', 'Disruptive', 'Driven', 'Fearless', 'Free', 'Gutsy', 'Independent', 'Inspired', 'Motivated', 'Rebellious', 'Revolutionary', 'Unafraid', 'Unconventional'],
        'Blue': ['Accomplish', 'Achieve', 'Affect', 'Assert', 'Cause', 'Command', 'Determine', 'Direct', 'Dominate', 'Drive', 'Empower', 'Establish', 'Guide', 'Impact', 'Impress', 'Influence', 'Inspire', 'Lead', 'Outpace', 'Outshine', 'Realize', 'Shape', 'Succeed', 'Transform', 'Win', 'Accomplished', 'Assertive', 'Authoritative', 'Commanding', 'Confident', 'Decisive', 'Distinguished', 'Dominant', 'Elite', 'Eminent', 'Established', 'Exceptional', 'Expert', 'First-class', 'First-rate', 'Impressive', 'Influential', 'Leading', 'Magnetic', 'Managerial', 'Masterful', 'Noble', 'Premier', 'Prestigious', 'Prominent', 'Proud', 'Strong'],
        'Yellow': ['Accelerate', 'Advance', 'Change', 'Conceive', 'Create', 'Engineer', 'Envision', 'Experiment', 'Dream', 'Ignite', 'Illuminate', 'Imagine', 'Innovate', 'Inspire', 'Invent', 'Pioneer', 'Progress', 'Shape', 'Spark', 'Solve', 'Transform', 'Unleash', 'Unlock', 'Advanced', 'Brilliant', 'Conceptual', 'Enterprising', 'Expert', 'Extraordinary', 'Forward-looking', 'Forward-thinking', 'Fresh', 'Future-minded', 'Future-thinking', 'Ingenious', 'Intelligent', 'Inventive', 'Leading-edge', 'Luminous', 'New', 'Pioneering', 'Reforming', 'Rising', 'Transformative', 'Visionary', 'World-changing', 'World-class'],
        'Green': ['Analyze', 'Discover', 'Examine', 'Expand', 'Explore', 'Extend', 'Inquire', 'Journey', 'Launch', 'Move', 'Pioneer', 'Pursue', 'Question', 'Reach', 'Search', 'Uncover', 'Venture', 'Wonder', 'Adventurous', 'Analytical', 'Curious', 'Discerning', 'Experiential', 'Exploratory', 'Fearless', 'Inquisitive', 'Intriguing', 'Investigative', 'Journeying', 'Mysterious', 'Philosophical', 'Pioneering', 'Questioning', 'Unbound', 'Unexpected'],
        'Purple': ['Accommodate', 'Assist', 'Befriend', 'Care', 'Collaborate', 'Connect', 'Embrace', 'Empower', 'Encourage', 'Foster', 'Give', 'Help', 'Nourish', 'Nurture', 'Promote', 'Protect', 'Provide', 'Serve', 'Share', 'Shepherd', 'Steward', 'Tend', 'Uplift', 'Value', 'Welcome', 'Affectionate', 'Attentive', 'Beneficial', 'Benevolent', 'Big-hearted', 'Caring', 'Charitable', 'Compassionate', 'Considerate', 'Encouraging', 'Friendly', 'Generous', 'Gentle', 'Helpful', 'Hospitable', 'Inclusive', 'Kind-hearted', 'Merciful', 'Missional', 'Neighborly', 'Nurturing', 'Protective', 'Responsible', 'Selfless', 'Supportive', 'Sympathetic', 'Thoughtful', 'Uplifting', 'Vocational', 'Warm'],
        'Maroon': ['Accomplish', 'Achieve', 'Build', 'Challenge', 'Commit', 'Compete', 'Contend', 'Dedicate', 'Defend', 'Devote', 'Drive', 'Endeavor', 'Entrust', 'Endure', 'Fight', 'Grapple', 'Grow', 'Improve', 'Increase', 'Overcome', 'Persevere', 'Persist', 'Press on', 'Pursue', 'Resolve', 'Tackle', 'Ambitious', 'Brave', 'Committed', 'Competitive', 'Consistent', 'Constant', 'Continuous', 'Courageous', 'Dedicated', 'Determined', 'Earnest', 'Industrious', 'Loyal', 'Persevering', 'Persistent', 'Proud', 'Purposeful', 'Relentless', 'Reliable', 'Resilient', 'Resolute', 'Steadfast', 'Strong', 'Tenacious', 'Tireless', 'Tough'],
        'Orange': ['Compose', 'Conceptualize', 'Conceive', 'Craft', 'Create', 'Design', 'Dream', 'Envision', 'Express', 'Fashion', 'Form', 'Imagine', 'Interpret', 'Make', 'Originate', 'Paint', 'Perform', 'Portray', 'Realize', 'Shape', 'Abstract', 'Artistic', 'Avant-garde', 'Colorful', 'Conceptual', 'Contemporary', 'Creative', 'Decorative', 'Eccentric', 'Eclectic', 'Evocative', 'Expressive', 'Imaginative', 'Interpretive', 'Offbeat', 'One-of-a-kind', 'Original', 'Uncommon', 'Unconventional', 'Unexpected', 'Unique', 'Vibrant', 'Whimsical'],
        'Pink': ['Arise', 'Aspire', 'Detail', 'Dream', 'Elevate', 'Enchant', 'Enrich', 'Envision', 'Exceed', 'Excel', 'Experience', 'Improve', 'Idealize', 'Imagine', 'Inspire', 'Perfect', 'Poise', 'Polish', 'Prepare', 'Refine', 'Uplift', 'Affectionate', 'Admirable', 'Age-less', 'Beautiful', 'Classic', 'Desirable', 'Detailed', 'Dreamy', 'Elegant', 'Enchanting', 'Enriching', 'Ethereal', 'Excellent', 'Exceptional', 'Experiential', 'Exquisite', 'Glamorous', 'Graceful', 'Idealistic', 'Inspiring', 'Lofty', 'Mysterious', 'Ordered', 'Perfect', 'Poised', 'Polished', 'Pristine', 'Pure', 'Refined', 'Romantic', 'Sophisticated', 'Spiritual', 'Timeless', 'Traditional', 'Virtuous', 'Visionary']
    }
    user_content = st.text_area('Paste your content here:')
    
    if st.button('Analyze'):
        st.session_state.init_done = True
        color_counts = analyze_text(user_content, color_keywords)
        st.session_state.updated_color_counts = color_counts.copy()
        
        st.session_state.initial_fig = draw_donut_chart(color_counts)
        st.subheader('Initial Donut Chart')
        st.plotly_chart(st.session_state.initial_fig)
        
        st.session_state.tone_scores = analyze_tone_with_gpt3(user_content, openai_api_key)
        
        sentences = re.split(r'[.!?]', user_content)
        st.session_state.sentence_to_colors = {}
        for sentence in sentences:
            if sentence.strip():
                initial_colors = [color for color, keywords in color_keywords.items() if any(keyword.lower() in sentence.lower() for keyword in keywords)]
                st.session_state.sentence_to_colors[sentence] = initial_colors
    
    if st.session_state.init_done:
        for tone in st.session_state.tone_scores.keys():
            st.session_state.tone_scores[tone] = st.slider(f"{tone}", 0, 10, int(st.session_state.tone_scores[tone]))
        
        st.session_state.tone_fig = go.Figure(data=[go.Bar(x=list(st.session_state.tone_scores.keys()), y=list(st.session_state.tone_scores.values()))])
        st.session_state.tone_fig.update_layout(xaxis_title='Tone', yaxis_title='Level')
        st.subheader("Updated Tone Analysis")
        st.plotly_chart(st.session_state.tone_fig)
        
        updated_color_counts = Counter()
        for sentence, initial_colors in st.session_state.sentence_to_colors.items():
            selected_colors = st.multiselect(f"{sentence}.", list(color_keywords.keys()), default=initial_colors)
            for color in selected_colors:
                updated_color_counts[color] += 1
        
        st.session_state.updated_fig = draw_donut_chart(updated_color_counts)
        st.subheader('Updated Donut Chart')
        st.plotly_chart(st.session_state.updated_fig)
        
        word_file_path = generate_word_doc(updated_color_counts, user_content, st.session_state.tone_scores, st.session_state.initial_fig, st.session_state.tone_fig, st.session_state.updated_fig)
        download_link = get_word_file_download_link(word_file_path, "Color_Personality_Analysis_Report.docx")
        st.markdown(download_link, unsafe_allow_html=True)

if __name__ == '__main__':
    main()
